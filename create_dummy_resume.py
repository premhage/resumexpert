from docx import Document

def create_dummy_resume():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer | Python Developer')
    
    doc.add_heading('Experience', level=1)
    p = doc.add_paragraph()
    p.add_run('Senior Python Developer').bold = True
    p.add_run('\nTech Corp | 2020 - Present').italic = True
    p.add_run('\n- Built scalable REST APIs using Flask and Django.')
    p.add_run('\n- Optimized database queries in PostgreSQL.')
    p.add_run('\n- Deployed applications using Docker and Kubernetes.')

    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, SQL, Docker, Kubernetes, AWS, Flask, Git, React')

    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.Sc. Computer Science | University of Tech')

    doc.save('dummy_resume.docx')
    print("Created dummy_resume.docx")

if __name__ == "__main__":
    create_dummy_resume()
